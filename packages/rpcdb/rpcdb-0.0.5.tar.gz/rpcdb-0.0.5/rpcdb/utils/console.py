import os
import time
from datetime import datetime
import calendar
import difflib
import sys
from fnmatch import fnmatch
from collections import deque
import argparse
import os
import codecs
import docx
import shutil
from contextlib import contextmanager
import inspect
# chat server
from socket import *
from datetime import datetime
from threading import Thread, Lock
# Chat client

sys.path.append(os.path.dirname(os.path.abspath(__file__)))
import chatclient

text = """
from setuptools import setup, find_packages

setup(
    name='',
    version='0.0',
    description='',
    long_description='',
    url='https://',
    author='',
    author_email='',
    license='GPL-3',
    platforms=['Windows 7','Windows 8', 'Windows 10','Unix','Mac OSX'],
    keywords ='',
    packages=find_packages(exclude=['docs*']),
    install_requires=[],
    )

"""

chatsever_running = False
def ChatServerLoop(host="", port=27000):
    global chatsever_running

    if not chatsever_running:
        clients = {}
        CHATS   = []

        s=socket(AF_INET, SOCK_DGRAM)
        s.bind((host, port))
        s.setblocking(1)
        chatsever_running = True
        print(f'\nRunning chat server on: {host} port {port}')

        while True:
            try:
                data, addr= s.recvfrom(1024)
                data = data.decode('utf-8')

                if "OFFLINE" in data:
                    del clients[addr]

                    message = data.replace("OFFLINE", '').strip() + " has disconnected" +"\n\n"
                    message = message.encode('utf-8')

                    for client in clients:
                        s.sendto(message, client)

                elif addr not in clients:
                    # Notify peers
                    name = str(data).replace("USERNAME", '').strip()
                    message = "{} is online...\n\n".format(name)
                    message = message.encode('utf-8')
                    CHATS.append(message)

                    for client in clients:
                        s.sendto(message, client)

                    # Send welcome message to new client
                    msg = "Welcome {}\n\n".format(name)
                    msg = msg.encode('utf-8')
                    clients[addr] = name
                    s.sendto(msg, addr)

                    # Send old messages
                    for msg in CHATS:
                        s.sendto(msg, addr)

                else:
                    message = data.encode('utf-8')
                    CHATS.append(message)
                    for client in clients:
                        s.sendto(message, client)

            except KeyboardInterrupt:
                pass
            except ConnectionResetError as e:
                pass
            except Exception:
                pass

    else:
        print(f"Chat Server already running on {host}: {port}")


@contextmanager
def working_dir(dirname):
    orig_dir = os.getcwd()
    try:
        os.chdir(dirname)
        yield
    finally:
        os.chdir(orig_dir)


class Package:
    def __init__(self, name):
        self.name = name

    def make_package(self):
        if os.path.exists(self.name):
            shutil.rmtree(self.name, ignore_errors=True)

        os.makedirs(self.name, exist_ok=True)
        self.pkg_name = os.path.abspath(self.name)

        with working_dir(self.pkg_name):
            os.makedirs('tests', exist_ok=True)

            with open('setup.py', 'w'):
                pass
            with open('__main__.py', 'w'):
                pass
            with open('README.md', 'w'):
                pass
            with open('DESCRIPTION.rst', 'w'):
                pass
            with open('LICENCE.txt', 'w'):
                pass
            with open('MANIFEST.in', 'w'):
                pass
            with open('__init__.py', 'w'):
                pass

            if not os.path.exists(self.name):
                os.makedirs(self.name, exist_ok=True)

            with open(os.path.join(self.name, '__init__.py'), 'w'):
                pass

            with open(os.path.join(self.name, 'main.py'), 'w'):
                pass

        self.write_setup()
        self.write_tests()

    def write_tests(self):
        with open(os.path.join(self.pkg_name, 'tests', 'tests.py'), 'w') as f:
            f.write('import unittest\n\n')
            f.write('class TestModule(unittest.TestCase):\n\n')
            f.write('   def setUp(self):\n\t\tpass\n\n')
            f.write('   def tearDown(self):\n\t\tpass\n')
            f.write("\n\nif __name__ == '__main__':\n    unittest.main()")


    def write_setup(self):
        with open(os.path.join(self.pkg_name, 'setup.py'), 'w') as f:
            f.write(text)

class linehistory:
    def __init__(self, lines, histlen=5):
        self.lines = lines
        self.history = deque(maxlen=histlen)

    def __iter__(self):
        for lineno, line in enumerate(self.lines, 1):
            self.history.append((lineno, line))
            yield line

    def clear(self):
        self.history.clear()


def find(filename, pat):
    showfile = False
    try:
        with codecs.open(filename) as f:
            lines = linehistory(f)
            for line in lines:
                if pat in line:
                    showfile= True
                    for lineno, hline in lines.history:
                        print('{}:{}'.format(lineno, hline), end='')

        if showfile:
            print(f"********* {filename} ***********\n\n")
            
    except Exception as e:
        pass


def allfiles(topdir):
    for dirname, subdirs, filenames in os.walk(topdir):
    
        for file in filenames:
            fullname = os.path.join(topdir, dirname, file)
            ext = os.path.splitext(fullname)[1]
            if ext in ['.txt', '.ini', '.py', '.yml', '.pyw', 
                       '.docx','.csv', '.in', '.cnf', '.conf']:
                yield fullname


def modified_within(modtime, label, top):
    now = time.time()
    for path, dirs, files in os.walk(top):
        for name in files:
            fullpath = os.path.join(path, name)
            if os.path.exists(fullpath):
                mtime = os.path.getmtime(fullpath)
                if label =='-s':
                    if mtime > (now - modtime):
                        print(fullpath)
                elif label=='-m':
                    if mtime > (now - modtime*60):
                        print(fullpath)

                elif label =='-h':
                    if mtime > (now - modtime*60*60):
                        print(fullpath)            


def follow(filename):
    f= codecs.open(filename, 'r', encoding='utf-8')
    f.seek(0, os.SEEK_END)

    while True:
        line = f.readline().rstrip()

        if not line:
            time.sleep(0.1)
            continue # Retry

        yield line   # Emit a line




class Bash:
    code = ""

    def walk(self):
        '''Walks recursively in current directory and lists
        all files with their sizes and modification dates
        in human readable format'''

        print("{:>30s} {:>20s} {:>30s}".format("File Name", "Size", "Mod Date"))

        for dirname, subdirs, files in os.walk(os.getcwd()):
            for file in files:
                size= os.path.getsize(file)
                moddate = time.ctime(os.path.getmtime(file))

                BYTES = size
                KB = size/1024
                MB = size/1024/1024
                GB = size/1024/1024/1024

                if size < 1024:
                    fsize= str(round(size, 2)) + " BYTES"
                elif size<1024*1024:
                    fsize = str(round(size/1024, 2)) + ' KB'
                elif size<1024*1024*1024:
                    fsize = str(round(size/(1024*1024), 2)) + ' MB'
                elif size<1024*1024*1024*1024:
                    fsize = str(round(size/(1024*1024*1024), 2)) + ' GB'

                print("{:>30s} {:>20s} {:>30s}".format(file, fsize, moddate))

    def modtime(self, *args):
        '''
        Lists all the files whose last modification time
        is specified. 
        Use as:
        >> modtime 20 -s C:/Users (files modified in 20 seconds)
        >> modtime 20 -m C:/Users (files modified in 20 minutes)
        >> modtime 20 -h C:/Users (files modified in 20 hours)
        >> modtime 20 -h . (files modified in 20 hours, in current directory)

        '''
        try:
            assert len(args) == 3, 'Usage: modtime [-s|-m|-h] directory\ntime_boundary=times since last file edit'
            t = float(args[0])
        except Exception as e:
            print(str(e))
            raise SystemExit(1)

        if '-s' in args:
            print('Looking for files modified within: %s seconds'%args[0])
            modified_within(float(args[0]), '-s', args[2])
        elif '-m' in args:
            print('Looking for files modified within: %s minutes'%args[0])
            modified_within(float(args[0]),'-m', args[2])
        elif '-h' in args:
            print('Looking for files modified within: %s hours'%args[0])
            modified_within(float(args[0]),'-h', args[2])
        else:
            print('Usage: modtime [-s|-m|-h] directory')
            raise SystemExit(1)

    def grep(self, pattern, dirname=None):
        '''Search for certain text recursively with pattern'''
        if dirname == None:
            dirname = os.path.abspath(".")

        for filename in allfiles(dirname):
            find(filename, pattern)

    def touch(self, *args):
        '''Makes an empty file in current directory
        Creates multiple files if multile filenames are
        supplied'''

        for filename in args:
            if not os.path.exists(filename):
                with open(filename, 'w'):
                    pass
                
    def pwd(self):
        '''Prints the current working directory'''
        print(os.getcwd())

    def clear(self):
        '''Clear the screen'''
        os.system('cls')
        
    def dir(self, args=None):
        ''' window's style listing of directory contents'''
        if args:
            os.system('dir %s'%args)
        else:
            os.system('dir')

    def ls(self):
        '''List the contents of a current directory and sub directories'''
        self.walk()

    def mkdir(self, dirname):
        '''Creates a new directory tree'''
        os.makedirs(dirname)

    def cd(self, dirname):
        '''Change the directory'''
        if dirname == '..':
            os.chdir(os.path.abspath(os.path.dirname(os.getcwd())))
        else:
            os.chdir(dirname)

    def diff(self, file1, file2):
        '''Prints the diferrence between two files'''

        if self.exists(file1):
            if self.exists(file2):
                fromlines = codecs.open(file1).readlines()
                tolines = codecs.open(file2).readlines()
                diff = difflib.context_diff(fromlines, tolines, file1, file2)
                sys.stdout.writelines(diff)
            else:
                print("File not Found: ", file2)
        else:
            print("File not Found: ", file1)

    def reset(self):
        '''Clears the screen'''
        os.system('cls')

    def find(self, pat, dirname=None):
        """Prints a table of files/file sizes matching a pattern
        from the specified directory"""
        if not dirname: 
            dirname= os.path.abspath('.')

        pattern = "*%s"%pat
        for dirname, sub, files in os.walk(dirname):
            for file in files:
                name = os.path.join(dirname, file)

                if fnmatch(name, pattern):
                    size= os.path.getsize(name)
                    BYTES = size
                    KB = size/1024
                    MB = size/1024/1024
                    GB = size/1024/1024/1024

                    if size < 1024:
                        print("{:>50}|{:>10.2f} BYTES".format(name, size))

                    elif size<1024*1024:
                        print("{:>50}|{:>10.2f} KB".format(name, size/1024))

                    elif size<1024*1024*1024:
                        print("{:>50}|{:>10.2f} MB".format(name, size/(1024*1024)))

                    elif size<1024*1024*1024*1024:
                        print("{:>50}|{:>10.2f} GB".format(name, size/(1024*1024*1024)))


    def exists(self, filename):
        return os.path.exists(os.path.abspath(filename))

    def tail(self, filename):
        if self.exists(filename):
            for line in follow(filename):
                print(line)
        else:
            print('FileNotFound: %s'%filename)

    def less(self, filename):
        with codecs.open(filename, 'rb') as f:
            print(f.read().decode('utf-8'))

    def exit(self):
        '''Quits the program'''
        sys.exit(0)

    def datetime(self):
        '''Prints the current date and time'''
        print(datetime.now().strftime('%d/%m/%Y %I:%M:%S %p'))

    def date(self):
        '''Prints the current date'''
        print(datetime.now().date().strftime(' %A, %d/%m/%Y'))

    def time(self):
        '''Prints the current time'''
        print(datetime.now().time().strftime('%I:%M:%S %p'))

    def cal(self, year=None):
        '''Displays a full calendar of the year'''
        cal = calendar.TextCalendar()
        cal.firstweekday = 6
        if year:
            print(cal.formatyear(int(year)))
        else:
            d = datetime.now().date()
            print(cal.formatyear(d.year))

    def __getattr__(self, name):
        return self.__dict__.get(name, None)

    def document(self, title, dirname='.'):
        """
        Creates a word document by recursively combing all files
        in the specified directory. Only for text(.txt) and python files
        (.py, .pyw)
        """
        def joinAllDocuments(dirname):
            for dirname, subdir, filenames in os.walk(dirname):
                for f in filenames:
                    ext = os.path.splitext(f)[1]
                    if ext in ('.py', '.txt', '.pyw'):
                        with open(f, 'r') as outfile:
                            self.code += "\n\n"+ f + "\n----------------\n"
                            data = outfile.read() +"\n\n"
                            self.code +=data

        if os.path.exists(dirname):
            joinAllDocuments(dirname)

            if self.code:
                d = docx.Document()
                d.add_heading(title, level=0)
                d.add_paragraph(self.code)
                d.save(title + '.docx')
            self.code = ""

    def pkg(self, name):
        '''
        Creates a python package directory
        with specified name in the current directory
        '''
        p = Package(name)
        p.make_package()
        p.write_setup()

    def rdel(self, ext, directory=None):
        if not directory:
            directory = os.path.abspath(".")
        '''Recursively delete all files matching extension ext'''
        if os.path.exists(os.path.abspath(directory)):
            ans = input("Delete all files "
                f"in {os.path.abspath(directory)} with extention {ext} ? [YES|NO] >> ")
            
            if ans != 'YES':
                raise SystemExit(0)

            for dirname, subdir, filenames in os.walk(directory):
                for f in filenames:
                    if f.endswith(ext):
                        print("Deleting: ", f)
                        os.remove(f)
                        time.sleep(0.2)

            print('Done!')
        else:
            print('No directory found!')

    def reboot(self):
        '''Reboots the system'''
        os.system("shutdown /r /t 0")

    def shutdown(self):
        '''Shuts down the computer'''
        os.system("shutdown /t 0")
            
    def help(self):
        '''Prints the built-in help of the Bash class with MRO'''
        print(help(self))
        print('*** End ****')

    def code(self, func):
        '''Prints the source code of a given function'''
        f= getattr(self, func)
        if callable(f):
            print(inspect.getsource(f))

    def doc(self, func):
        '''Prints the documentation of a given function'''
        f= getattr(self, func)
        if callable(f):
            print(f.__doc__)

    def extract(self, archive_name):
        '''Unpacks an archive'''
        f= os.path.abspath(archive_name)
        fname = os.path.basename(f).split(".")[0]
        dirname = os.path.join(os.path.dirname(f), fname)
        os.makedirs(dirname, exist_ok=True)
        with working_dir(dirname):
            shutil.unpack_archive(f)
            
    def compress(self, archive_name, fmt, input_folder):
        '''Make an arhive/compressed file
        Specify output '''
        
        try:
            shutil.make_archive("./%s"%archive_name, fmt, input_folder)
        except Exception as e:
            print(e)
            avail = shutil.get_archive_formats()
            print("The following the available archive formats")
            for n in avail:
                print(n)

    def chatserver(self, host="", port=17000):
        '''Run a chat server on a given host and port'''
        thread = Thread(target=ChatServerLoop, args = (host, int(port)), daemon=True)
        thread.start()

    def chatclient(self, host, port):
        '''Run a chat client on a given host and port'''
        t = Thread(target=chatclient.main, args=(host, int(port)), daemon=True)
        t.start()


def run():
    bash = Bash()
    print("Console with Unix Utilities. Developed by Dr Abiira Nathan")

    while True:
        try:
            path = os.getcwd()
            cmdIn = input("\n(shell)%s> "%path)

            if not cmdIn:
                continue

            cmd, *args = cmdIn.strip().split()
            func = getattr(bash, cmd)
            args = tuple(arg.replace('_', ' ') for arg in args)

            if func:
                try:
                    if args:
                        func(*args)
                    else:
                        func()
                except Exception as e:
                    print(e)

            else:
                # Try running in command line
                try:
                    os.system(cmdIn.rstrip())
                except Exception as e:
                    print(e)

        except KeyboardInterrupt:
            print('\nKeyboardInterrupt by user')
            raise SystemExit(0)

        except EOFError:
            raise SystemExit(0)


